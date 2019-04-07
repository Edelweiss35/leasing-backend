# -*- coding: utf-8 -*-

import os, re, getopt, sys
from os.path import splitext
from nltk.tokenize import sent_tokenize
from docx import Document
from docx.shared import Inches
from .classify import classify
from io import StringIO
from six import BytesIO as StringIO
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
import docx
import pdftotext
import docx2txt
import pdb
import unicodedata
import math
import string
from .fileReader import FileReader
from django.conf import settings


def convert_pdf(fname, pages=None):
    if not pages:
        pagenums = set()
    else:
        pagenums = set(pages)

    output = StringIO()
    manager = PDFResourceManager()
    converter = TextConverter(manager, output, laparams=LAParams())
    interpreter = PDFPageInterpreter(manager, converter)

    infile = open(fname, 'rb')
    for page in PDFPage.get_pages(infile, pagenums):
        interpreter.process_page(page)
    infile.close()
    converter.close()
    text = output.getvalue()
    output.close
    return text

def convert_pdf2text(fname):
    with open(settings.BASE_DIR + '/leasingai/ai/temp/' + fname, "rb") as f:
        pdf = pdftotext.PDF(f)
    text = "\n\n".join(pdf)
    return text

def convert_docx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def convert_doc2txt(fname):
    print("convert_doc2txt BEGINS")
    text = docx2txt.process(fname)
    print("text converted is %s \n" % text[:25])
    return(text)

def clean(text):
    """
    Remove any extra whitespace and line breaks as needed.
    """
    # print("text in clean is", text)
    # Replace linebreaks with spaces
    text = text.replace("\n", " ").replace("\r", " ").replace("\t", " ")

    # Remove any leeding or trailing whitespace
    text = text.strip()

    # Remove consecutive spaces
    # text = re.sub(" +", " ", text)
    
    text = text.replace(u'\u2019', u'\'').encode('ascii', 'ignore').decode('utf-8')
    text = re.sub('[^\x00-\x7F]+',' ', text)

    text

    return ''.join([i if ord(i) < 128 else ' ' for i in text])


def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )



def remove_clause_number(text):
    for i in range(len(text)):
        if(text[i].isalpha()):
            break
    return text[i:]



def create_table(filename, clauses):
    complete_path = settings.BASE_DIR + '/leasingai/ai/temp/' + filename
    reader = FileReader()
    print(complete_path) 
    l=[]
    name = splitext(filename)[0]
    document = Document()
    document.add_heading('Clause Extraction', 0)
    table = document.add_table(rows=1, cols=6)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Amended No.'
    hdr_cells[1].text = 'Clause No.'
    hdr_cells[2].text = 'ClauseName'
    hdr_cells[3].text = 'Lease required amendment'
    hdr_cells[4].text = 'Required action'
    hdr_cells[5].text = 'Lessor response to amendment Agreed/Not agreed'
    index = 0
    total_index = 1
    # 4 Clauses to Extract
    for clause in clauses:
        t={}
        
        t['index'] = total_index
        t['action'] = clause['action']
        t['clausename'] = clause['clause_name']
        t['reason'] = clause['reason']
        t['agree']=""
        to_find = [clause['clause_name']]
        row_cells = table.add_row().cells
        row_cells[0].text = str(t['index'])
        row_cells[1].text = ""
        row_cells[2].text = t['clausename']
        row_cells[3].text = t['reason']
        row_cells[5].text = t['agree']
        row_cells[4].text = t['action']
        total_index += 1
        if(clause['text']):
            cleaned_string = ''.join(c for c in clause['text'] if valid_xml_char_ordinal(c))
            print("_______________________ cleasned string is _________________", cleaned_string)
            f = remove_clause_number(cleaned_string)
            index += 1
            t['keep'] = 0
            clause_no=""
            # lst = {}
            try:
                temp_no = 0
                clause_no = re.match(r"([0-9]+\.[0-9]+\.[0-9]+|[0-9]+\.[0-9]+|[0-9]+\.)", cleaned_string)
                if(clause_no != None):
                    clause_no = clause_no.group()
                    temp_no = int(clause_no.replace('.',''))
                else:
                    clause_no = ""
                lst = reader.find_with(clause['text'], [clause['key_text']], headers_only=False)
                if(len(lst) > 0 and len(lst[0]) > 0):
                    try:
                        if(string.ascii_letters.find(lst[0].replace("(", "").replace(')','').replace('.','')[:-1]) == 0):
                          clause_no = clause_no + ','.join(lst)
                        elif(int(lst[0].replace("(", "").replace(')','').replace('.','')) != temp_no):
                          clause_no = clause_no + ','.join(lst)
                        else:
                          clause_no = ','.join(lst)
                    except ValueError:
                        clause_no = clause_no + ','.join(lst)
            except Exception as e:
                print (e)
                # raise Exception
                print(complete_path)
                result = reader.find(path=complete_path, to_find=to_find, headers_only=True)
                print (result)
                raise Exception
                try:
                    clause_no = result[0]
                except:
                    clause_no = ""
                print(clause_no)
            if( clause_no is not None and clause_no != ""):
                t['clause_no'] = clause_no
                row_cells[1].text = t['clause_no']
                l.append(t)
            else:
                l.append(t)
        else:
            t['clause_no'] = ""
            t['keep'] = 1
            l.append(t)
    document.add_page_break()
    document.save(settings.BASE_DIR + '/leasingai/ai/temp/'+name+'.docx')
    return l


def make_doc(text):
    l = []
    document = Document()
    document.add_heading('Clause Extraction', 0)
    table = document.add_table(rows=1, cols=1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Index'

    index = 0
    if(text):
        t={}
        cleaned_string = ''.join(c for c in text if valid_xml_char_ordinal(c))
        f = remove_clause_number(cleaned_string)
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        row_cells[1].text = re.match(r"([0-9]+\.[0-9]+|[0-9]+\.)", cleaned_string).group()
        row_cells[2].text = f
        row_cells[3].text = clause['clauseaction']
        index += 1
        t['index'] = index
        t['clause_no'] = re.match(r"([0-9]+\.[0-9]+|[0-9]+\.)", cleaned_string).group()

        t['action'] = clause['clauseaction']
        t['keep'] = 0
        t['data'] = f
        t['reason'] = clause['clausereason']
        l.append(t)

    document.add_page_break()

    document.save(settings.BASE_DIR + '/leasingai/ai/temp/'+name+'.docx')
    return l  


def main_create_table(data, name="random"):
    l=[]

    document = Document()
    document.add_heading('Clause Extraction', 0)
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Amended No.'
    hdr_cells[1].text = 'Clause No.'
    hdr_cells[2].text = 'ClauseName'
    hdr_cells[3].text = 'Lease required amendment'
    hdr_cells[4].text = 'Required action'
    hdr_cells[5].text = 'Lessor response to amendment Agreed/Not agreed'
    for i in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(data[i]['index'])
        row_cells[1].text = str(data[i]['clause_no'])
        row_cells[2].text = str(data[i]['data'])
        row_cells[3].text = str(data[i]['action'])
        row_cells[4].text = str(data[i]['reason'])   
    document.add_page_break()

    document.save(settings.BASE_DIR + '/leasingai/ai/temp/'+name+'.docx')

# ---------------- get clause content ------------

def get_clausecontent(clause_name, text):
    print (clause_name)
    clause_regex = "((([0-9]+\.[0-9]+|[0-9]+\.)([0-9]+\.[0-9]+|[0-9]+\.)) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_name, clause_name, clause_name,clause_name, clause_name, clause_name, clause_name)
    clause_match = re.search(clause_regex, text)
    return clause_match

# ---------------- get clause content(alpha) -----------

def get_alphacontent(clause_name, text):
    print("----------------------------hi ==========")
    print (clause_name)
    # text.replace("(","1111")
    # text.replace(")", "2222")
    print(text)
    clause_regex_with_alpha = "((([a-z]|[A-Z]+)(\.\s)|(\([A-Z]|[a-z]\)\s|\([A-Z]|[a-z]\)+))(%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_name, clause_name, clause_name,clause_name, clause_name, clause_name, clause_name)
    clause_match = re.search(clause_regex_with_alpha, text)
    if(clause_match == None):
      clr_text = re.sub(" +", " ", text)
      clause_match = re.search(clause_regex_with_alpha, clr_text)
      if(clause_match == None):
        clause_regex = "((%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_name, clause_name, clause_name,clause_name, clause_name, clause_name, clause_name)
        clause_match = re.search(clause_regex, text)

    return clause_match


def main(filename, clauses):

    if(splitext(filename)[1] == '.pdf'):
        print("File extension is pdf and is verified by main function")
        text = convert_pdf2text(filename)
    elif(splitext(filename)[1] == '.docx'):
        print("File extension is docx and is verified by main function")
        text = convert_doc2txt(filename)
    else:
        return
    clean_text = clean(text)
    clean_text = clean(clean_text)

    temp_clause = []

    for clause in clauses:
        print("<<<<<<<<<<<<<<<<<<<<<< clause in for loop is >>>>>>>>>>>>>>>>> ", clause)
        while True:
            try:
                clause['clause_name'] = [clean(_clause['text']) for _clause in clause['clause_name']]
                clause_names = clause['clause_name']
                print("______________________ clause names are ___________________", clause_names)
                for clause_name in clause_names:
                    clause_text = ""
                    if(clause_name == "GST"):
                        clause_name = "Goods and Services Tax"
                        clause_nameU = clause_name.upper()
                        clause_regex2 = "(([0-9]{,9}\.+[0-9]{,9}\.+[0-9]+|[0-9]{,9}\.+[0-9]{,9}\.+|[0-9]{,9}\.+) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.[0-9]+\.[0-9]|[0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_name,
                                                                                                                                                                                                                                                       clause_name,
                                                                                                                                                                                                                                                       clause_name,
                                                                                                                                                                                                                                                       clause_name,
                                                                                                                                                                                                                                                       clause_name,
                                                                                                                                                                                                                                                       clause_name,
                                                                                                                                                                                                                                                       clause_name)
                        clause_regex21 = "(([0-9]+\.[0-9]+\.[0-9]+|[0-9]+\.[0-9]\.+|[0-9]+\.) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]{2})" % (clause_name,
                                                                                                                                                                                                                 clause_name,
                                                                                                                                                                                                                 clause_name,
                                                                                                                                                                                                                 clause_name,
                                                                                                                                                                                                                 clause_name,
                                                                                                                                                                                                                 clause_name,
                                                                                                                                                                                                                 clause_name)
                        clause_regexU = "(([0-9]+\.[0-9]+|[0-9]+\.) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_nameU,
                                                                                                                                                                                    clause_nameU,
                                                                                                                                                                                    clause_nameU,
                                                                                                                                                                                    clause_nameU,
                                                                                                                                                                                    clause_nameU,
                                                                                                                                                                                    clause_nameU,
                                                                                                                                                                                    clause_nameU)
                        clause_regex4 = "(([0-9]+\.[0-9]+|[0-9]+\.) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)(\.\s{2})([a-z]|[A-Z])" % (clause_name,
                                                                                                                                                                           clause_name,
                                                                                                                                                                           clause_name,
                                                                                                                                                                           clause_name,
                                                                                                                                                                           clause_name,
                                                                                                                                                                           clause_name,
                                                                                                                                                                           clause_name)
                        if clause_name == "Goods and Services Tax":
                            clause_nameG = "GST"
                            clause_regexG = "(([0-9]+\.[0-9]+|[0-9]+\.) (%s +|%s\. +|%s\: +|%s\; +|%s\ . +|%s\ : +|%s\ ; +)(\(|[a-z]|[A-Z])(.|\n)+?)([0-9]+\.|[0-9]+\.[0-9]+) [A-Z]" % (clause_nameG,
                                                                                                                                                                                        clause_nameG,
                                                                                                                                                                                        clause_nameG,
                                                                                                                                                                                        clause_nameG,
                                                                                                                                                                                        clause_nameG,
                                                                                                                                                                                        clause_nameG,
                                                                                                                                                                                        clause_nameG)
                            clause_matchG = re.search(clause_regexG, clean_text)
                            if(clause_matchG):
                                clause_text = clause_matchG.groups()[0]
                        clause_match = re.search(clause_regex2, clean_text)
                        clause_match21 = re.search(clause_regex21, clean_text)
                        clause_matchU = re.search(clause_regexU, clean_text)
                        clause_match4 = re.search(clause_regex4, clean_text)
                        if(clause_match):
                            clause_text = clause_match.groups()[0]
                        if(clause_match21):
                            clause_text = clause_match21.groups()[0]
                        if(clause_matchU):
                            clause_text = clause_matchU.groups()[0]
                        if(clause_match4 and clause_text != "" and len(clause_text) < 30):
                            clause_text = clause_match4.groups()[0]
                        clause_text = re.sub(" +", " ", clause_text)
                        if(clause_text == ""):
                            clause_match1 = get_clausecontent(clause_name,
                                                              clean_text)
                            if clause_match1:
                                clause_text = clause_match1.groups()[0]
                        print("____________clause_text_______++++",
                              clause_text)
                        if(clause_text == ""):
                            clause_match2 = get_alphacontent(clause_name,
                                                             clean_text)
                            print("\n\n\nclause_match2 \n\n\n", clause_match2)
                            if clause_match2:
                                clause_text = clause_match2.groups()[0]
                                last_index = len(clause_match2.groups()) - 1
                                print("\n clause_text is %s and last_index is %s \n" % (clause_text, last_index))
                                print("\n clause_match2.groups()\n",
                                      clause_match2.groups())
                                try:
                                    proceeding_number_clause = re.match('[0-9]{1,2}\.[0-9]{1,2}|[0-9]{1,2}\.|[0-9]{2,3}\.', clause_match2.groups()[last_index]).group()
                                except AttributeError:
                                    proceeding_number_clause = None
                                print("procedding number clause is %s " % proceeding_number_clause)
                                if proceeding_number_clause:
                                    try:
                                        int_clause_number = int(proceeding_number_clause)
                                        clause_text = str(int_clause_number - 1 ) + ". " + clause_text
                                    except Exception:
                                        float_clause_number = float(proceeding_number_clause)
                                        frac, whole = math.modf(float_clause_number)
                                        if(round(frac, 1) > 0.1 and round(frac, 1) <= 0.9 ):
                                            clause_text = str(float_clause_number - 0.1 ) + ". " + clause_text
                                        else:
                                            previous_paras = sent_tokenize(clean_text.split(clause_text)[0])
                                            previous_paras.reverse()
                                            for sentence in previous_paras:
                                                print("_______________________ sentence is _________________________________", sentence)
                                                sentence = sentence.replace('(','').replace(')','')
                                                sent_number = re.match('[0-9]{1,2}\.[0-9]{1,2}', sentence)
                                                if sent_number:
                                                    clause_text = sent_number.group() + ". " + clause_text
                                                    break
                        if(clause_text == "" or clause_text is None):
                            clause_text = clean_text
                            print("clause in cond 1 is", clause)
                            clause['text'] = clause_text
                            print("clause in cond 1.1 is", clause)
                            # temp_clause.append(clause.copy())
                            break
                        else:
                            # clean_text = clean_text.replace(clause_text, '')
                            print("_______________________ clean text is _________________________________", clean_text)
                            clean_text = re.sub(" +", " ", clean_text)
                            clean_text = clean_text.replace(re.sub(r'^[0-9.]+\ ', '', clause_text), '')
                            print("clause in cond 2 is", clause)
                            clause['text'] = clause_text
                            temp_clause.append(clause.copy())
                            # temp_clause.append(clause)
                            # clauses.append(clause)
                            if(len(temp_clause) > 20):
                                break
                            continue
            except Exception as e:
                print("\n Exceptoin is %s \n" % e)
                break
    # Create The Table
    name = splitext(filename)[0]
    # k = create_table(name, clauses)
    # print(clean_text)
    return temp_clause
